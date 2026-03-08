import os
import sys
import customtkinter as ctk
from tkinter import messagebox
import docx
import datetime as dt


def get_base_dir():
    if getattr(sys, 'frozen', False):
        return os.path.dirname(sys.executable)
    return os.path.dirname(os.path.abspath(__file__))


class InvoiceAutomation:
    def __init__(self):
        self.root = ctk.CTk()
        self.root.title("Invoice Automation")
        self.root.geometry("600x450")
        self.root.resizable(False, False)
        self.build_ui()

    def build_ui(self):
        ctk.CTkLabel(self.root, text="Invoice Automation Tool").pack(pady=5)

        ctk.CTkLabel(self.root, text="Customer Name:").pack(pady=5)
        self.name = ctk.CTkEntry(self.root, width=250)
        self.name.pack(pady=5)

        ctk.CTkLabel(self.root, text="Invoice Number:").pack(pady=5)
        self.invoice = ctk.CTkEntry(self.root, width=250)
        self.invoice.pack(pady=5)

        ctk.CTkLabel(self.root, text="Service:").pack(pady=5)
        self.service = ctk.CTkEntry(self.root, width=250)
        self.service.pack(pady=5)

        ctk.CTkLabel(self.root, text="Amount:").pack(pady=5)
        self.amount = ctk.CTkEntry(self.root, width=250)
        self.amount.pack(pady=5)

        ctk.CTkButton(self.root, text="Create Invoice", command=self.create_invoice).pack(pady=5)

    @staticmethod
    def replace_text(paragraph, old_text, new_text):
        if old_text in paragraph.text:
            paragraph.text = paragraph.text.replace(old_text, new_text)

    def create_invoice(self):
        base_dir = get_base_dir()
        template_path = os.path.join(base_dir, 'template.docx')
        output_path = os.path.join(base_dir, 'filled.docx')  # Save next to exe/script

        if not os.path.exists(template_path):
            messagebox.showerror("Error", f"template.docx not found in:\n{base_dir}")
            return

        doc = docx.Document(template_path)

        replacements = {
            "[Date]": dt.datetime.today().strftime("%Y-%m-%d"),
            "[customer_name]": self.name.get(),
            "[invoice_number]": self.invoice.get(),
            "[service]": self.service.get(),
            "[amount]": self.amount.get()
        }

        for paragraph in doc.paragraphs:
            for old_text, new_text in replacements.items():
                self.replace_text(paragraph, old_text, new_text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for old_text, new_text in replacements.items():
                            self.replace_text(paragraph, old_text, new_text)

        doc.save(output_path)
        messagebox.showinfo("Success", f"Invoice saved to:\n{output_path}")


if __name__ == "__main__":
    app = InvoiceAutomation()
    app.root.mainloop()